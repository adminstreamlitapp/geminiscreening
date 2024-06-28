import os
import docx
import pandas as pd
from PyPDF2 import PdfReader
import streamlit as st
import re

# Model configuration for text generation
generation_config = {
    "temperature": 0.4,
    "top_p": 1,
    "top_k": 32,
    "max_output_tokens": 4096,
}

# Define safety settings for content generation
safety_settings = [
    {"category": f"HARM_CATEGORY_{category}", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
    for category in ["HARASSMENT", "HATE_SPEECH", "SEXUALLY_EXPLICIT", "DANGEROUS_CONTENT"]
]

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_file(file_path):
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        return ""

def generate_response_from_gemini(input_text, prompt):
    # This is a placeholder for the Gemini AI response.
    return "This is a placeholder for the Gemini AI response."

def normalize_text(text):
    return re.sub(r'\s+', '', text.lower())

def evaluate_resume(resume_text, required_skills, optional_skills):
    normalized_resume_text = normalize_text(resume_text)
    
    def create_patterns(skill):
        normalized_skill = normalize_text(skill)
        patterns = [
            normalized_skill,
            re.escape(skill.lower().replace(' ', '')),  # Handle potential spaces between words
        ]
        return patterns

    def match_skills(skills, text):
        found_skills = {skill: False for skill in skills}
        for skill in skills:
            patterns = create_patterns(skill)
            for pattern in patterns:
                if re.search(pattern, text):
                    found_skills[skill] = True
                    break
        return found_skills

    required_skills_list = required_skills.lower().split('\n')
    optional_skills_list = optional_skills.lower().split('\n')
    
    required_skills_found = match_skills(required_skills_list, normalized_resume_text)
    optional_skills_found = match_skills(optional_skills_list, normalized_resume_text)
    
    required_match_percentage = sum(required_skills_found.values()) / len(required_skills_found) * 100
    optional_match_percentage = sum(optional_skills_found.values()) / len(optional_skills_found) * 100
    
    return required_skills_found, optional_skills_found, required_match_percentage, optional_match_percentage

st.title('Resume Screening')

# Columns for required and optional skills
col1, col2 = st.columns(2)

with col1:
    required_skills = st.text_area("Enter Required Skills (one skill per line):", height=200)
with col2:
    optional_skills = st.text_area("Enter Optional Skills (one skill per line):", height=200)

resumes_directory = st.text_input("Enter directory containing resumes:")

if st.button("Analyze Resumes"):
    if resumes_directory and required_skills and optional_skills:
        results = []
        for resume_file_name in os.listdir(resumes_directory):
            if not resume_file_name.startswith('~$'):
                resume_path = os.path.join(resumes_directory, resume_file_name)
                if os.path.isfile(resume_path) and (resume_file_name.endswith('.docx') or resume_file_name.endswith('.pdf')):
                    resume_text = extract_text_from_file(resume_path)
                    required_skills_found, optional_skills_found, required_match_percentage, optional_match_percentage = evaluate_resume(resume_text, required_skills, optional_skills)
                    
                    # Generate experience and education summaries using Gemini AI
                    experience_summary = generate_response_from_gemini(resume_text, "Extract the experience summary from the following resume:")
                    education_summary = generate_response_from_gemini(resume_text, "Extract the education summary from the following resume:")
                    
                    results.append({
                        "Resume name": resume_file_name,
                        "Required Skills": ", ".join([skill for skill, found in required_skills_found.items() if found]),
                        "Optional Skills": ", ".join([skill for skill, found in optional_skills_found.items() if found]),
                        "Required skills % match": f"{required_match_percentage:.2f}%",
                        "Optional skills % match": f"{optional_match_percentage:.2f}%",
                        "Experience Summary": experience_summary,
                        "Education Summary": education_summary,
                    })
        
        results_df = pd.DataFrame(results)
        st.dataframe(results_df)
    else:
        st.warning("Please enter required skills, optional skills, and resumes directory")
